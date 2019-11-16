#!usr/bin/python
# -*-coding:utf-8-*-


# 查看了网上的各种方法，大部分还是接受将doc文件强制另存为docx文件（使用代码转换，而不是直接修改后缀），在读取即可，需要另外安装win32com模块，注意就是直接使用 pip install win32com 安装不成功，需要用
# python -m pip install pypiwin32

import docx
import win32com.client as wc


#TODO 软件是否安装影响本部分程序功能，Try...还是有软件要求。
# doc文件另存为docx
word = wc.Dispatch("Word.Application")
#wps文件另存为docx
wps = wc.Dispatch('wps.application')
kwps = wc.Dispatch('kwps.application')
doc = word.Documents.Open(r"G:\\OARemainder\\通知文件下载\\关于做好2019年度综合绩效考核相关工作的通知关于做好2019年度综合绩效考核相关工作的通知.doc")
# 上面的地方只能使用完整绝对地址，相对地址找不到文件，且，只能用“\\”，不能用“/”，哪怕加了 r 也不行，涉及到将反斜杠看成转义字符。
doc.SaveAs(r"G:\\OARemainder\\通知文件下载\\关于做好2019年度综合绩效考核相关工作的通知关于做好2019年度综合绩效考核相关工作的通知.docx", 12, False, "", True, "",
           False, False, False,
           False)  # 转换后的文件,12代表转换后为docx文件
# doc.SaveAs(r"F:\\***\\***\\appendDoc\\***.docx", 12)#或直接简写
# 注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
doc.Close
word.Quit

# path = "appendDoc/***.docx"
# file = docx.Document(path)
# for p in file.paragraphs:
#     print(p.text)

# 没有具体实现了解，只当看了下，需要安装库 pdfminer 使用


#coding=utf-8
from docx import Document
import os,sys

def search_word(filename,word):
    #打开文档
    document = Document(filename)
    # document = Document(r'C:\Users\Cheng\Desktop\kword\words\wind.docx')
    print filename
    #读取每段资料
    l = [ paragraph.text.encode('gb2312') for paragraph in document.paragraphs];
    #输出并观察结果，也可以通过其他手段处理文本即可
    for i in l:
        i=i.strip()
        # print i
        if i.find(word)!=-1:
            print filename, i

def get_process_files(root_dir):
    """process all files in directory"""
    cur_dir=os.path.abspath(root_dir)
    file_list=os.listdir(cur_dir)
    process_list=[]
    for file in file_list:
        fullfile=cur_dir+"\\"+file
        if os.path.isfile(fullfile):
            process_list.append(fullfile)
        elif os.path.isdir(fullfile):
            dir_extra_list=get_process_files(fullfile)
            if len(dir_extra_list)!=0:
                for x in dir_extra_list:
                    process_list.append(x)
    return process_list

def find_files(root_dir,word):
    process_list=get_process_files(root_dir)
    for files in process_list:
        search_word(files, word)

if __name__=='__main__':
    #文件根目录
    root_dir=sys.argv[1]
    #要搜索的关键字
    word=sys.argv[2]
    try:
        find_files(root_dir,word)
    except:
        pass